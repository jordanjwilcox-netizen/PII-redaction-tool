import os
import docx
from engine import redact


def process_file(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        result = redact(content)

        new_path = path.rsplit(".", 1)[0] + "_redacted.txt"
        with open(new_path, "w", encoding="utf-8") as f:
            f.write(result["text"])

        return result

    if ext == ".docx":
        return redact_docx(path)

    return {"text": "", "entities": []}


def process_folder(folder):
    files = []

    for root_dir, _, filenames in os.walk(folder):
        for f in filenames:
            if f.lower().endswith((".txt", ".docx")):
                files.append(os.path.join(root_dir, f))

    for file_path in files:
        process_file(file_path)

    return len(files)


def redact_docx(path):
    doc = docx.Document(path)

    full_text = extract_full_docx_text(doc)
    result = redact(full_text)

    entities = result.get("entities", [])

    apply_entities_to_docx(doc, entities)

    new_path = path.rsplit(".", 1)[0] + "_redacted.docx"
    doc.save(new_path)

    return {
        "text": f"Saved DOCX: {new_path}",
        "entities": entities
    }


def extract_full_docx_text(doc):
    parts = []

    for p in doc.paragraphs:
        parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)

    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)

        for p in section.footer.paragraphs:
            parts.append(p.text)

    return "\n".join(parts)


def apply_entities_to_docx(doc, entities):
    for p in doc.paragraphs:
        apply_entities_to_paragraph(p, entities)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_entities_to_paragraph(p, entities)

    for section in doc.sections:
        for p in section.header.paragraphs:
            apply_entities_to_paragraph(p, entities)

        for p in section.footer.paragraphs:
            apply_entities_to_paragraph(p, entities)


def apply_entities_to_paragraph(paragraph, entities):
    if not paragraph.text.strip():
        return

    original = paragraph.text
    redacted = original

    for entity in entities:
        entity_type = entity.get("type", "REDACTED")
        value = entity.get("value", "")

        if value:
            redacted = redacted.replace(value, f"[{entity_type}]")

    if redacted == original:
        return

    # Clear normal runs
    for run in paragraph.runs:
        run.text = ""

    # Clear visible text inside Word hyperlink XML
    for hyperlink in paragraph._p.findall(".//w:hyperlink", paragraph._p.nsmap):
        for r in hyperlink.findall(".//w:r", paragraph._p.nsmap):
            for t in r.findall(".//w:t", paragraph._p.nsmap):
                t.text = ""

    # Put final redacted paragraph text back once
    if paragraph.runs:
        paragraph.runs[0].text = redacted
    else:
        paragraph.add_run(redacted)